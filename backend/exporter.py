"""
Module d'exportation d'ebooks en différents formats
Supporte : PDF, EPUB, MOBI, DOCX, HTML
"""

from reportlab.lib import colors
from reportlab.lib.pagesizes import A4, letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak, Table, TableStyle
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY, TA_LEFT
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont

from ebooklib import epub
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

import markdown2
from io import BytesIO
import os
import re


class EbookExporter:
    """Classe pour exporter les ebooks dans différents formats"""
    
    def __init__(self, ebook_data):
        """
        Initialize with ebook data from MongoDB
        
        Args:
            ebook_data: dict containing ebook information
        """
        self.ebook = ebook_data
        self.title = ebook_data.get('title', 'Sans titre')
        self.author = ebook_data.get('author', 'Anonyme')
        self.chapters = ebook_data.get('chapters', [])
        self.description = ebook_data.get('description', '')
        self.cover = ebook_data.get('cover', {})
        self.toc = ebook_data.get('toc', [])
        self.legal_pages = ebook_data.get('legal_pages', {})
    
    def export_to_pdf(self) -> BytesIO:
        """
        Export ebook to PDF format with professional layout
        
        Returns:
            BytesIO: PDF file in memory
        """
        buffer = BytesIO()
        doc = SimpleDocTemplate(
            buffer,
            pagesize=A4,
            rightMargin=72,
            leftMargin=72,
            topMargin=72,
            bottomMargin=18,
        )
        
        # Styles
        styles = getSampleStyleSheet()
        
        # Custom styles
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=24,
            textColor=colors.HexColor('#3B82F6'),
            spaceAfter=30,
            alignment=TA_CENTER,
            fontName='Helvetica-Bold'
        )
        
        author_style = ParagraphStyle(
            'AuthorStyle',
            parent=styles['Normal'],
            fontSize=14,
            textColor=colors.HexColor('#6B7280'),
            spaceAfter=12,
            alignment=TA_CENTER,
            fontName='Helvetica'
        )
        
        chapter_title_style = ParagraphStyle(
            'ChapterTitle',
            parent=styles['Heading1'],
            fontSize=18,
            textColor=colors.HexColor('#8B5CF6'),
            spaceAfter=20,
            spaceBefore=20,
            fontName='Helvetica-Bold'
        )
        
        body_style = ParagraphStyle(
            'BodyText',
            parent=styles['Normal'],
            fontSize=11,
            leading=16,
            alignment=TA_JUSTIFY,
            spaceAfter=12,
            fontName='Helvetica'
        )
        
        # Build document
        story = []
        
        # Cover page
        story.append(Spacer(1, 2*inch))
        story.append(Paragraph(self.title, title_style))
        story.append(Spacer(1, 0.5*inch))
        story.append(Paragraph(f"par {self.author}", author_style))
        
        # Add tagline if exists
        if self.cover.get('tagline'):
            tagline_style = ParagraphStyle(
                'Tagline',
                parent=styles['Normal'],
                fontSize=12,
                textColor=colors.HexColor('#F97316'),
                spaceAfter=12,
                alignment=TA_CENTER,
                fontName='Helvetica-Oblique'
            )
            story.append(Spacer(1, 1*inch))
            story.append(Paragraph(self.cover['tagline'], tagline_style))
        
        story.append(PageBreak())
        
        # Legal pages if available
        if self.legal_pages:
            # Copyright page
            if self.legal_pages.get('copyright_page'):
                copyright_style = ParagraphStyle(
                    'Copyright',
                    parent=styles['Normal'],
                    fontSize=10,
                    alignment=TA_CENTER,
                    spaceAfter=8,
                    fontName='Helvetica'
                )
                copyright_lines = self.legal_pages['copyright_page'].split('\n')
                for line in copyright_lines:
                    if line.strip():
                        story.append(Paragraph(line.strip(), copyright_style))
                        story.append(Spacer(1, 4))
                story.append(PageBreak())
            
            # Legal mentions
            if self.legal_pages.get('legal_mentions'):
                story.append(Paragraph("Mentions Légales", chapter_title_style))
                story.append(Spacer(1, 0.2*inch))
                legal_lines = self.legal_pages['legal_mentions'].split('\n')
                for line in legal_lines:
                    if line.strip():
                        story.append(Paragraph(line.strip(), body_style))
                        story.append(Spacer(1, 6))
                story.append(PageBreak())
        
        # Table of contents (enhanced with subtitles)
        story.append(Paragraph("Table des Matières", chapter_title_style))
        story.append(Spacer(1, 0.3*inch))
        
        toc_style = ParagraphStyle(
            'TOCEntry',
            parent=styles['Normal'],
            fontSize=11,
            spaceAfter=8,
            fontName='Helvetica-Bold'
        )
        
        toc_subtitle_style = ParagraphStyle(
            'TOCSubtitle',
            parent=styles['Normal'],
            fontSize=9,
            textColor=colors.HexColor('#6B7280'),
            spaceAfter=4,
            leftIndent=20,
            fontName='Helvetica'
        )
        
        for idx, chapter in enumerate(self.chapters):
            # Main chapter entry
            chapter_type = chapter.get('type', 'chapter')
            if chapter_type == 'introduction':
                toc_entry = f"Introduction"
            elif chapter_type == 'conclusion':
                toc_entry = f"Conclusion"
            else:
                toc_entry = f"Chapitre {chapter['number']} : {chapter['title']}"
            
            story.append(Paragraph(toc_entry, toc_style))
            
            # Add subtitles if available from TOC data
            if self.toc and idx < len(self.toc):
                subtitles = self.toc[idx].get('subtitles', [])
                for subtitle in subtitles:
                    story.append(Paragraph(f"• {subtitle}", toc_subtitle_style))
            
            story.append(Spacer(1, 6))
        
        story.append(PageBreak())
        
        # Chapters
        for chapter in self.chapters:
            # Chapter title
            chapter_heading = f"Chapitre {chapter['number']}: {chapter['title']}"
            if chapter.get('type') == 'introduction':
                chapter_heading = chapter['title']
            elif chapter.get('type') == 'conclusion':
                chapter_heading = chapter['title']
            
            story.append(Paragraph(chapter_heading, chapter_title_style))
            story.append(Spacer(1, 0.2*inch))
            
            # Chapter content - split by paragraphs
            content = chapter.get('content', '')
            paragraphs = content.split('\n\n')
            
            for para in paragraphs:
                if para.strip():
                    # Check if it's a section with 🔹 marker
                    if para.strip().startswith('🔹'):
                        subtitle = para.strip().replace('🔹', '').strip()
                        subtitle_style = ParagraphStyle(
                            'SubtitleStyle',
                            parent=styles['Heading2'],
                            fontSize=14,
                            textColor=colors.HexColor('#8B5CF6'),
                            spaceAfter=10,
                            spaceBefore=15,
                            fontName='Helvetica-Bold'
                        )
                        story.append(Paragraph(subtitle, subtitle_style))
                    # Fallback: still handle ## format if present (for backward compatibility)
                    elif para.strip().startswith('##'):
                        subtitle = para.strip().replace('##', '').strip().replace('#', '').strip()
                        subtitle_style = ParagraphStyle(
                            'SubtitleStyle',
                            parent=styles['Heading2'],
                            fontSize=14,
                            textColor=colors.HexColor('#8B5CF6'),
                            spaceAfter=10,
                            spaceBefore=15,
                            fontName='Helvetica-Bold'
                        )
                        story.append(Paragraph(subtitle, subtitle_style))
                    # Remove any stray # symbols at the start
                    elif para.strip().startswith('#'):
                        clean_text = para.strip().lstrip('#').strip()
                        # If it looks like a heading, treat it as subtitle
                        if len(clean_text) < 100 and not '. ' in clean_text[:20]:
                            subtitle_style = ParagraphStyle(
                                'SubtitleStyle',
                                parent=styles['Heading2'],
                                fontSize=14,
                                textColor=colors.HexColor('#8B5CF6'),
                                spaceAfter=10,
                                spaceBefore=15,
                                fontName='Helvetica-Bold'
                            )
                            story.append(Paragraph(clean_text, subtitle_style))
                        else:
                            # Regular paragraph
                            clean_para = clean_text.replace('\n', ' ')
                            story.append(Paragraph(clean_para, body_style))
                            story.append(Spacer(1, 6))
                    else:
                        # Clean and add paragraph - remove any stray markdown symbols
                        clean_para = para.strip().replace('\n', ' ')
                        # Remove markdown bold/italic markers if present
                        clean_para = re.sub(r'\*\*(.*?)\*\*', r'<b>\1</b>', clean_para)
                        clean_para = re.sub(r'\*(.*?)\*', r'<i>\1</i>', clean_para)
                        story.append(Paragraph(clean_para, body_style))
                        story.append(Spacer(1, 6))
            
            story.append(PageBreak())
        
        # Build PDF
        doc.build(story)
        buffer.seek(0)
        return buffer
    
    def export_to_epub(self) -> BytesIO:
        """
        Export ebook to EPUB format (compatible with e-readers)
        
        Returns:
            BytesIO: EPUB file in memory
        """
        book = epub.EpubBook()
        
        # Set metadata
        book.set_identifier(self.ebook.get('_id', 'id123456'))
        book.set_title(self.title)
        book.set_language('fr')
        book.add_author(self.author)
        
        if self.description:
            book.add_metadata('DC', 'description', self.description)
        
        # Add ISBN if available
        if self.legal_pages.get('isbn') and self.legal_pages['isbn'] != 'Non attribué':
            book.add_metadata('DC', 'identifier', self.legal_pages['isbn'], {'id': 'isbn'})
        
        # Create legal pages chapter if available
        epub_chapters = []
        
        if self.legal_pages:
            legal_content = "<h1>Informations Légales</h1>"
            
            if self.legal_pages.get('copyright_page'):
                legal_content += "<div style='text-align: center; margin: 2em 0;'>"
                copyright_lines = self.legal_pages['copyright_page'].split('\n')
                for line in copyright_lines:
                    if line.strip():
                        legal_content += f"<p>{line.strip()}</p>"
                legal_content += "</div>"
            
            if self.legal_pages.get('legal_mentions'):
                legal_content += "<h2>Mentions Légales</h2>"
                legal_lines = self.legal_pages['legal_mentions'].split('\n')
                for line in legal_lines:
                    if line.strip():
                        legal_content += f"<p>{line.strip()}</p>"
            
            legal_chapter = epub.EpubHtml(
                title='Informations Légales',
                file_name='legal.xhtml',
                lang='fr'
            )
            legal_chapter.content = legal_content
            book.add_item(legal_chapter)
            epub_chapters.append(legal_chapter)
        
        # Create chapters
        epub_chapters = []
        
        for idx, chapter in enumerate(self.chapters):
            chapter_title = f"{chapter['number']}. {chapter['title']}"
            if chapter.get('type') == 'introduction':
                chapter_title = chapter['title']
            elif chapter.get('type') == 'conclusion':
                chapter_title = chapter['title']
            
            # Create chapter content with HTML
            content = f"<h1>{chapter_title}</h1>"
            
            # Convert content to HTML paragraphs
            text_content = chapter.get('content', '')
            paragraphs = text_content.split('\n\n')
            
            for para in paragraphs:
                if para.strip():
                    # Handle 🔹 sections
                    if para.strip().startswith('🔹'):
                        subtitle = para.strip().replace('🔹', '').strip()
                        content += f"<h2>{subtitle}</h2>"
                    # Fallback for ## format
                    elif para.strip().startswith('##'):
                        subtitle = para.strip().replace('##', '').strip().replace('#', '').strip()
                        content += f"<h2>{subtitle}</h2>"
                    # Clean stray # symbols
                    elif para.strip().startswith('#'):
                        clean_text = para.strip().lstrip('#').strip()
                        if len(clean_text) < 100 and not '. ' in clean_text[:20]:
                            content += f"<h2>{clean_text}</h2>"
                        else:
                            content += f"<p>{clean_text}</p>"
                    else:
                        clean_para = para.strip().replace('\n', ' ')
                        # Convert markdown bold/italic to HTML
                        clean_para = re.sub(r'\*\*(.*?)\*\*', r'<strong>\1</strong>', clean_para)
                        clean_para = re.sub(r'\*(.*?)\*', r'<em>\1</em>', clean_para)
                        content += f"<p>{clean_para}</p>"
            
            # Create EPUB chapter
            epub_chapter = epub.EpubHtml(
                title=chapter['title'],
                file_name=f'chap_{idx+1}.xhtml',
                lang='fr'
            )
            epub_chapter.content = content
            book.add_item(epub_chapter)
            epub_chapters.append(epub_chapter)
        
        # Define Table Of Contents
        book.toc = epub_chapters
        
        # Add default NCX and Nav file
        book.add_item(epub.EpubNcx())
        book.add_item(epub.EpubNav())
        
        # Define CSS style
        style = '''
        @namespace epub "http://www.idpf.org/2007/ops";
        body {
            font-family: Cambria, Liberation Serif, Bitstream Vera Serif, Georgia, Times, Times New Roman, serif;
            line-height: 1.6;
            margin: 5%;
        }
        h1 {
            text-align: left;
            font-weight: 200;
            color: #3B82F6;
            margin-bottom: 1em;
        }
        h2 {
            color: #8B5CF6;
            margin-top: 1.5em;
            margin-bottom: 0.5em;
        }
        p {
            text-align: justify;
            margin-bottom: 1em;
        }
        '''
        
        nav_css = epub.EpubItem(
            uid="style_nav",
            file_name="style/nav.css",
            media_type="text/css",
            content=style
        )
        book.add_item(nav_css)
        
        # Create spine
        book.spine = ['nav'] + epub_chapters
        
        # Write to BytesIO
        buffer = BytesIO()
        epub.write_epub(buffer, book, {})
        buffer.seek(0)
        return buffer
    
    def export_to_docx(self) -> BytesIO:
        """
        Export ebook to DOCX format (editable Word document)
        
        Returns:
            BytesIO: DOCX file in memory
        """
        doc = Document()
        
        # Add title
        title = doc.add_heading(self.title, 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title.runs[0].font.color.rgb = RGBColor(59, 130, 246)
        
        # Add author
        author_para = doc.add_paragraph(f"par {self.author}")
        author_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        author_para.runs[0].font.size = Pt(14)
        author_para.runs[0].font.color.rgb = RGBColor(107, 114, 128)
        
        # Add tagline if exists
        if self.cover.get('tagline'):
            tagline_para = doc.add_paragraph(self.cover['tagline'])
            tagline_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            tagline_para.runs[0].font.italic = True
            tagline_para.runs[0].font.color.rgb = RGBColor(249, 115, 22)
        
        doc.add_page_break()
        
        # Table of contents
        doc.add_heading('Table des Matières', level=1)
        for chapter in self.chapters:
            toc_entry = f"{chapter['number']}. {chapter['title']}"
            doc.add_paragraph(toc_entry, style='List Number')
        
        doc.add_page_break()
        
        # Chapters
        for chapter in self.chapters:
            # Chapter title
            chapter_heading = f"Chapitre {chapter['number']}: {chapter['title']}"
            if chapter.get('type') == 'introduction':
                chapter_heading = chapter['title']
            elif chapter.get('type') == 'conclusion':
                chapter_heading = chapter['title']
            
            heading = doc.add_heading(chapter_heading, level=1)
            heading.runs[0].font.color.rgb = RGBColor(139, 92, 246)
            
            # Chapter content
            content = chapter.get('content', '')
            paragraphs = content.split('\n\n')
            
            for para in paragraphs:
                if para.strip():
                    if para.strip().startswith('##'):
                        # Subtitle
                        subtitle = para.strip().replace('##', '').strip()
                        sub_heading = doc.add_heading(subtitle, level=2)
                        sub_heading.runs[0].font.color.rgb = RGBColor(30, 64, 175)
                    else:
                        # Regular paragraph
                        p = doc.add_paragraph(para.strip())
                        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            
            doc.add_page_break()
        
        # Save to BytesIO
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer
    
    def export_to_html_flipbook(self) -> BytesIO:
        """
        Export ebook to HTML format (interactive flipbook)
        
        Returns:
            BytesIO: HTML file in memory (zipped with assets)
        """
        html_content = f"""<!DOCTYPE html>
<html lang="fr">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{self.title}</title>
    <style>
        * {{
            margin: 0;
            padding: 0;
            box-sizing: border-box;
        }}
        
        body {{
            font-family: 'Georgia', serif;
            background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
            min-height: 100vh;
            padding: 20px;
        }}
        
        .flipbook-container {{
            max-width: 1200px;
            margin: 0 auto;
            background: white;
            border-radius: 20px;
            box-shadow: 0 20px 60px rgba(0,0,0,0.3);
            overflow: hidden;
        }}
        
        .cover {{
            background: linear-gradient(135deg, #3B82F6 0%, #8B5CF6 100%);
            color: white;
            padding: 100px 50px;
            text-align: center;
            min-height: 600px;
            display: flex;
            flex-direction: column;
            justify-content: center;
        }}
        
        .cover h1 {{
            font-size: 3em;
            margin-bottom: 20px;
            text-shadow: 2px 2px 4px rgba(0,0,0,0.2);
        }}
        
        .cover .author {{
            font-size: 1.5em;
            opacity: 0.9;
            margin-bottom: 30px;
        }}
        
        .cover .tagline {{
            font-size: 1.2em;
            font-style: italic;
            color: #FCD34D;
            margin-top: 40px;
        }}
        
        .page {{
            padding: 60px;
            min-height: 600px;
            display: none;
        }}
        
        .page.active {{
            display: block;
        }}
        
        .page h1 {{
            color: #8B5CF6;
            font-size: 2.5em;
            margin-bottom: 30px;
            border-bottom: 3px solid #3B82F6;
            padding-bottom: 15px;
        }}
        
        .page h2 {{
            color: #1E40AF;
            font-size: 1.8em;
            margin-top: 30px;
            margin-bottom: 15px;
        }}
        
        .page p {{
            line-height: 1.8;
            text-align: justify;
            margin-bottom: 20px;
            color: #374151;
            font-size: 1.1em;
        }}
        
        .navigation {{
            background: #F3F4F6;
            padding: 20px;
            display: flex;
            justify-content: space-between;
            align-items: center;
            border-top: 1px solid #E5E7EB;
        }}
        
        .nav-button {{
            background: linear-gradient(135deg, #3B82F6, #8B5CF6);
            color: white;
            border: none;
            padding: 12px 30px;
            border-radius: 8px;
            font-size: 1em;
            cursor: pointer;
            transition: transform 0.2s, box-shadow 0.2s;
        }}
        
        .nav-button:hover {{
            transform: translateY(-2px);
            box-shadow: 0 10px 20px rgba(59, 130, 246, 0.3);
        }}
        
        .nav-button:disabled {{
            opacity: 0.5;
            cursor: not-allowed;
        }}
        
        .page-info {{
            color: #6B7280;
            font-size: 0.9em;
        }}
        
        .toc {{
            padding: 60px;
        }}
        
        .toc h1 {{
            color: #8B5CF6;
            margin-bottom: 30px;
        }}
        
        .toc-item {{
            padding: 15px;
            margin: 10px 0;
            background: #F9FAFB;
            border-left: 4px solid #3B82F6;
            cursor: pointer;
            transition: all 0.3s;
        }}
        
        .toc-item:hover {{
            background: #EEF2FF;
            transform: translateX(10px);
        }}
        
        .toc-item .chapter-number {{
            color: #8B5CF6;
            font-weight: bold;
            margin-right: 10px;
        }}
    </style>
</head>
<body>
    <div class="flipbook-container">
        <!-- Cover Page -->
        <div class="page active" id="page-0">
            <div class="cover">
                <h1>{self.title}</h1>
                <div class="author">par {self.author}</div>
                {f'<div class="tagline">{self.cover.get("tagline", "")}</div>' if self.cover.get('tagline') else ''}
            </div>
        </div>
        
        <!-- Table of Contents -->
        <div class="page" id="page-1">
            <div class="toc">
                <h1>Table des Matières</h1>
"""
        
        for idx, chapter in enumerate(self.chapters):
            html_content += f"""
                <div class="toc-item" onclick="goToPage({idx + 2})">
                    <span class="chapter-number">{chapter['number']}.</span>
                    <span class="chapter-title">{chapter['title']}</span>
                </div>
"""
        
        html_content += """
            </div>
        </div>
        
"""
        
        # Add chapters
        for idx, chapter in enumerate(self.chapters):
            chapter_heading = f"Chapitre {chapter['number']}: {chapter['title']}"
            if chapter.get('type') == 'introduction':
                chapter_heading = chapter['title']
            elif chapter.get('type') == 'conclusion':
                chapter_heading = chapter['title']
            
            html_content += f"""
        <div class="page" id="page-{idx + 2}">
            <h1>{chapter_heading}</h1>
"""
            
            # Add content
            content = chapter.get('content', '')
            paragraphs = content.split('\n\n')
            
            for para in paragraphs:
                if para.strip():
                    if para.strip().startswith('##'):
                        subtitle = para.strip().replace('##', '').strip()
                        html_content += f"            <h2>{subtitle}</h2>\n"
                    else:
                        clean_para = para.strip().replace('\n', ' ')
                        html_content += f"            <p>{clean_para}</p>\n"
            
            html_content += """
        </div>
"""
        
        total_pages = len(self.chapters) + 2
        
        html_content += f"""
        
        <!-- Navigation -->
        <div class="navigation">
            <button class="nav-button" id="prevBtn" onclick="prevPage()">← Précédent</button>
            <div class="page-info">
                <span id="currentPage">1</span> / <span id="totalPages">{total_pages}</span>
            </div>
            <button class="nav-button" id="nextBtn" onclick="nextPage()">Suivant →</button>
        </div>
    </div>
    
    <script>
        let currentPage = 0;
        const totalPages = {total_pages};
        
        function showPage(pageNum) {{
            document.querySelectorAll('.page').forEach(page => {{
                page.classList.remove('active');
            }});
            
            document.getElementById('page-' + pageNum).classList.add('active');
            document.getElementById('currentPage').textContent = pageNum + 1;
            
            // Update button states
            document.getElementById('prevBtn').disabled = (pageNum === 0);
            document.getElementById('nextBtn').disabled = (pageNum === totalPages - 1);
            
            currentPage = pageNum;
            
            // Smooth scroll to top
            window.scrollTo({{ top: 0, behavior: 'smooth' }});
        }}
        
        function nextPage() {{
            if (currentPage < totalPages - 1) {{
                showPage(currentPage + 1);
            }}
        }}
        
        function prevPage() {{
            if (currentPage > 0) {{
                showPage(currentPage - 1);
            }}
        }}
        
        function goToPage(pageNum) {{
            showPage(pageNum);
        }}
        
        // Keyboard navigation
        document.addEventListener('keydown', (e) => {{
            if (e.key === 'ArrowRight') nextPage();
            if (e.key === 'ArrowLeft') prevPage();
        }});
        
        // Initialize
        showPage(0);
    </script>
</body>
</html>
"""
        
        buffer = BytesIO()
        buffer.write(html_content.encode('utf-8'))
        buffer.seek(0)
        return buffer
    
    def export_to_mobi(self) -> BytesIO:
        """
        Export ebook to MOBI format (Kindle)
        Note: MOBI is being phased out by Amazon in favor of EPUB
        This creates an EPUB that can be converted to MOBI
        
        Returns:
            BytesIO: EPUB file (to be converted to MOBI)
        """
        # MOBI format requires external tool (kindlegen or calibre)
        # For now, return EPUB format with note
        # In production, use subprocess to call ebook-convert
        return self.export_to_epub()
